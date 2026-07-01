import importlib.util
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "submission" / "generate_vdt_report.py"
spec = importlib.util.spec_from_file_location("generate_vdt_report", MODULE_PATH)
assert spec is not None
assert spec.loader is not None
generate_vdt_report = importlib.util.module_from_spec(spec)
spec.loader.exec_module(generate_vdt_report)


def test_emit_body_renders_markdown_table_as_word_table() -> None:
    document = Document()

    generate_vdt_report.emit_body(
        document,
        [
            "Kết quả HotpotQA pilot 200 truy vấn:",
            "",
            "| Phương pháp | Full-support@10 | Recall@10 | Độ trễ p95 |",
            "| --- | ---: | ---: | ---: |",
            "| `es_bm25` | 0.365 | 0.6025 | 359.63 ms |",
            "| `tv_hybrid` | 0.545 | 0.7500 | 2061-3089 ms |",
        ],
    )

    assert len(document.tables) == 1
    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == [
        "Phương pháp",
        "Full-support@10",
        "Recall@10",
        "Độ trễ p95",
    ]
    assert [cell.text for cell in table.rows[1].cells] == [
        "es_bm25",
        "0.365",
        "0.6025",
        "359.63 ms",
    ]
    assert [cell.text for cell in table.rows[2].cells] == [
        "tv_hybrid",
        "0.545",
        "0.7500",
        "2061-3089 ms",
    ]


def test_submission_markdown_contains_expanded_vdt_sections() -> None:
    source = (ROOT / "submission" / "bao-cao-vdt-2026.md").read_text(encoding="utf-8")

    required_phrases = [
        "TurboVec không thay Elasticsearch",
        "tv_hybrid",
        "tv_bridge_title_entities_rrf",
        "Ablation 200 truy vấn",
        "full test 7,405 truy vấn",
        "Pyserini",
        "MDR",
        "Beam Retrieval",
        "IRCoT",
    ]

    for phrase in required_phrases:
        assert phrase in source
