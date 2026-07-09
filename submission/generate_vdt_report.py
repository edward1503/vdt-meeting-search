from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "submission" / "bao-cao-vdt-2026.md"
OUTPUTS = [
    ROOT / "submission" / "bao-cao-vdt-2026.docx",
    ROOT / "submission" / "bao-cao-vdt-2026-theo-mau.docx",
]


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", text).strip()


def set_columns(section, count: int = 2, space_twips: str = "284") -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), space_twips)


def set_page(section) -> None:
    # Match the converted VDT template: A4, top 2cm, other margins 1.5cm.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(0.5)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.first_line_indent = Cm(0.5)
    normal.paragraph_format.line_spacing = Pt(14)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None, font: str = "Times New Roman"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_title(document: Document, title: str, author: str, mentor: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    add_run(p, title, bold=True, size=16, font="Times New Roman")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, author, italic=True, size=14, font="Times New Roman")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, mentor, italic=True, size=14, font="Times New Roman")


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, clean_inline(text), bold=True, size=13)


def add_body_paragraph(document: Document, text: str, *, compact: bool = False) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0 if compact else 0.5)
    if compact:
        p.paragraph_format.left_indent = Cm(0.35)
    add_run(p, clean_inline(text), size=12)


def parse_front_matter(lines: list[str]) -> tuple[str, str, str, list[str]]:
    nonempty = [(i, line.strip()) for i, line in enumerate(lines) if line.strip()]
    if len(nonempty) < 3:
        raise ValueError("Markdown source must contain title, author line, and mentor footnote.")
    title_i, title = nonempty[0]
    author_i, author = nonempty[1]
    mentor_i, mentor = nonempty[2]
    if title.startswith("# "):
        title = title[2:].strip()
    body_start = mentor_i + 1
    return clean_inline(title), clean_inline(author), clean_inline(mentor), lines[body_start:]


def table_rows(block: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in block:
        cells = [clean_inline(cell) for cell in line.strip().strip("|").split("|")]
        if not cells or all(re.fullmatch(r":?-+:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, margin_twips: str = "80") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), margin_twips)
        node.set(qn("w:type"), "dxa")

def write_table_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(11)
    add_run(paragraph, text, bold=bold, size=9.5)

def emit_table(document: Document, block: list[str]) -> None:
    rows = table_rows(block)
    if len(rows) <= 1:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.rows[row_index].cells[column_index]
            if row_index == 0:
                set_cell_shading(cell, "EDEDED")
            write_table_cell(
                cell,
                row[column_index] if column_index < len(row) else "",
                bold=row_index == 0,
            )

    spacer = document.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.space_after = Pt(2)


def emit_body(document: Document, lines: list[str]) -> None:
    paragraph: list[str] = []
    table: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            add_body_paragraph(document, " ".join(paragraph))
            paragraph = []

    def flush_table() -> None:
        nonlocal table
        if table:
            emit_table(document, table)
            table = []

    for raw in lines + [""]:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("|"):
            flush_paragraph()
            table.append(stripped)
            continue

        flush_table()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(document, stripped[3:].strip())
            continue

        paragraph.append(stripped)


def build_docx(output_path: Path) -> None:
    document = Document()
    style_document(document)
    set_page(document.sections[0])

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    title, author, mentor, body_lines = parse_front_matter(lines)
    add_title(document, title, author, mentor)

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    set_page(body_section)
    set_columns(body_section)
    emit_body(document, body_lines)
    final_section = document.add_section(WD_SECTION.CONTINUOUS)
    set_page(final_section)

    document.save(output_path)


def main() -> None:
    for output in OUTPUTS:
        build_docx(output)
        print(output)


if __name__ == "__main__":
    main()
